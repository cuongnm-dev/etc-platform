"""
One-off: expand TKKT + TKCS docx templates with new placeholders per NĐ 45/2026 Đ13.

Adds new Heading 2/3 + paragraph with {{ tkcs.FIELD or '[CẦN BỔ SUNG: ...]' }} style
placeholder. Appends at end of document (user can reorder in Word later).
"""
from __future__ import annotations
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
TMPL_DIR = ROOT / "src" / "etc_docgen" / "assets" / "templates"


def add_tkkt_fields(doc_path: Path):
    """Add business_overview + design_principles to TKKT template."""
    doc = Document(str(doc_path))
    # Append at end
    doc.add_heading("Tổng quan nghiệp vụ", level=2)
    doc.add_paragraph(
        "{{ architecture.business_overview or '[CẦN BỔ SUNG: Tổng quan nghiệp vụ — Business view per Khung KT CPĐT 4.0]' }}"
    )
    doc.add_heading("Nguyên tắc thiết kế kiến trúc", level=2)
    doc.add_paragraph(
        "{{ architecture.design_principles or '[CẦN BỔ SUNG: Nguyên tắc thiết kế — scalability, security-by-design, modularity, reusability]' }}"
    )
    doc.save(str(doc_path))
    print(f"TKKT: added 2 placeholders → {doc_path.name}")


def add_tkcs_fields(doc_path: Path):
    """Add 20 new fields to TKCS template per NĐ 45/2026 Đ13 full expansion."""
    doc = Document(str(doc_path))

    additions = [
        ("Mục tiêu đầu tư", "objectives",
         "Mục tiêu tổng quát, mục tiêu cụ thể với KPI đo được"),
        # Section 5 — technology analysis
        ("Pattern kiến trúc phần mềm", "software_arch_pattern",
         "Phân tích lựa chọn microservices / monolith / SOA"),
        ("Lựa chọn hệ quản trị CSDL", "dbms_choice",
         "Phân tích lựa chọn DBMS (RDBMS vs NoSQL, vendor)"),
        ("Lựa chọn hệ điều hành", "os_choice",
         "Phân tích lựa chọn HĐH máy chủ (Linux distro / Windows Server)"),
        ("Tiêu chuẩn áp dụng", "standards",
         "Các tiêu chuẩn TCVN, ISO, IEEE, OpenAPI áp dụng"),
        # Section 6 — detailed design
        ("Thiết kế phần mềm", "software_design",
         "Thiết kế phần mềm tổng quan — module, layer, interface"),
        ("Thiết kế hạ tầng", "infrastructure_design",
         "Thiết kế hạ tầng — server, network, storage"),
        # Section 7 — security
        ("Thiết kế bảo mật chi tiết", "security_design",
         "Thiết kế ATTT theo TCVN 11930 — 5 nhóm biện pháp (quản lý, kỹ thuật, vật lý, con người, vận hành)"),
        ("Giải pháp kỹ thuật ATTT", "security_tech",
         "WAF, SIEM, IDS/IPS, mã hóa end-to-end"),
        # Section 8 — operations prep
        ("Chuẩn bị nguồn lực vận hành", "prep_resources",
         "Nhân lực, thiết bị, hạ tầng chuẩn bị cho vận hành"),
        ("Chuẩn bị quy chế vận hành", "prep_policies",
         "Quy chế, quy trình, SOP vận hành"),
        ("Hỗ trợ và đào tạo người dùng", "user_support",
         "Phương án hỗ trợ, đào tạo, truyền thông tới người dùng"),
        # Section 9 — schedule
        ("Các mốc triển khai (milestones)", "milestones",
         "Các milestone chính: khởi động, thiết kế, triển khai, UAT, go-live"),
        ("Kế hoạch tiến độ chi tiết", "schedule",
         "Bảng tiến độ Gantt / WBS từng hạng mục"),
        # Section 10 — budget detail
        ("Dự toán chi tiết", "budget_detail",
         "Chi tiết dự toán theo TT 04/2020/TT-BTTTT — từng hạng mục"),
        ("Chi phí vận hành chi tiết", "opex",
         "Chi phí O&M hàng năm — nhân lực, license, hạ tầng, điện nước"),
        ("Bảo hành, bảo trì", "warranty",
         "Điều khoản bảo hành, thời hạn bảo trì sau go-live"),
        # Section 11 — PM detail
        ("Hình thức quản lý dự án", "pm_form",
         "BQL chuyên trách / kiêm nhiệm / thuê tư vấn QLDA"),
        ("Các bên liên quan", "stakeholders",
         "Chủ đầu tư, nhà thầu, tư vấn giám sát, đơn vị thẩm định"),
        ("Phương pháp quản lý dự án", "pm_method",
         "Waterfall / Agile / Hybrid — áp dụng cho từng giai đoạn"),
    ]

    for heading, field, hint in additions:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(
            "{{ tkcs." + field + " or '[CẦN BỔ SUNG: " + hint + "]' }}"
        )

    doc.save(str(doc_path))
    print(f"TKCS: added {len(additions)} placeholders → {doc_path.name}")


if __name__ == "__main__":
    add_tkkt_fields(TMPL_DIR / "thiet-ke-kien-truc.docx")
    add_tkcs_fields(TMPL_DIR / "thiet-ke-co-so.docx")
    # Also overlay copy at src/etc_docgen/templates/
    overlay = ROOT / "src" / "etc_docgen" / "templates" / "thiet-ke-kien-truc.docx"
    if overlay.exists():
        add_tkkt_fields(overlay)
    print("DONE.")
