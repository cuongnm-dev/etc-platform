#!/usr/bin/env python3
"""Build nghien-cuu-kha-thi.docx (NCKT Jinja2 template) programmatically.

Generates a DOCX skeleton matching outline data/outlines/nghien-cuu-kha-thi/nd45-2026.md
with {{ nckt.sections['X.Y'] }} placeholders consumed by docxtpl render.

Run:
  python scripts/build_nckt_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "src/etc_platform/assets/templates/nghien-cuu-kha-thi.docx"

# (number, title, level, key) — level: 1=H1 chương, 2=H2 mục, 3=H3 tiểu mục
# key=None means a parent heading without direct content placeholder
OUTLINE: list[tuple[str, str, int, str | None]] = [
    # 1
    ("1.", "TỔNG QUAN VỀ DỰ ÁN", 1, None),
    ("1.1.", "Thông tin chung về dự án", 2, "1.1"),
    ("1.2.", "Căn cứ pháp lý lập dự án", 2, "1.2"),
    ("1.3.", "Yêu cầu về an toàn thông tin của dự án", 2, None),
    ("1.3.1.", "Bảo vệ nội dung dự án", 3, "1.3.1"),
    ("1.3.2.", "Bảo vệ thông tin tài chính", 3, "1.3.2"),
    ("1.3.3.", "Bảo vệ giải pháp kỹ thuật", 3, "1.3.3"),
    ("1.3.4.", "Bảo vệ danh sách thiết bị", 3, "1.3.4"),
    ("1.3.5.", "Bảo vệ tài liệu và báo cáo", 3, "1.3.5"),
    ("1.3.6.", "Bảo vệ tiến độ dự án", 3, "1.3.6"),
    ("1.3.7.", "Giám sát và đánh giá tuân thủ", 3, "1.3.7"),
    # 2
    ("2.", "SỰ CẦN THIẾT ĐẦU TƯ", 1, None),
    ("2.1.", "Hiện trạng tổ chức và nghiệp vụ", 2, "2.1"),
    ("2.2.", "Hiện trạng ứng dụng công nghệ thông tin", 2, None),
    ("2.2.1.", "Về ứng dụng CNTT", 3, "2.2.1"),
    ("2.2.2.", "Về nhân lực CNTT", 3, "2.2.2"),
    ("2.3.", "Hiện trạng hạ tầng CNTT và các dự án đang triển khai", 2, None),
    ("2.3.1.", "Hiện trạng hạ tầng CNTT", 3, "2.3.1"),
    ("2.3.2.", "Các dự án CNTT đang triển khai", 3, "2.3.2"),
    ("2.3.3.", "Phân tích đánh giá hiện trạng ứng dụng CNTT", 3, "2.3.3"),
    ("2.3.4.", "Phân tích đánh giá hạ tầng và khả năng tận dụng", 3, "2.3.4"),
    ("2.4.", "Sự cần thiết đầu tư dự án", 2, "2.4"),
    # 3
    ("3.", "ĐÁNH GIÁ SỰ PHÙ HỢP VỚI QUY HOẠCH PHÁT TRIỂN KINH TẾ - XÃ HỘI", 1, None),
    ("3.1.", "Sự tuân thủ Kiến trúc Chính phủ điện tử", 2, "3.1"),
    ("3.2.", "Phù hợp Quy hoạch ứng dụng, phát triển CNTT của ngành", 2, "3.2"),
    # 4
    ("4.", "PHÂN TÍCH, XÁC ĐỊNH MỤC TIÊU, NHIỆM VỤ, KẾT QUẢ ĐẦU RA CỦA DỰ ÁN", 1, None),
    ("4.1.", "Mục tiêu đầu tư", 2, None),
    ("4.1.1.", "Mục tiêu tổng quát", 3, "4.1.1"),
    ("4.1.2.", "Mục tiêu cụ thể", 3, "4.1.2"),
    ("4.2.", "Quy mô đầu tư", 2, "4.2"),
    ("4.3.", "Thời gian thực hiện dự án", 2, "4.3"),
    ("4.4.", "Lựa chọn hình thức đầu tư", 2, "4.4"),
    # 5
    ("5.", "PHÂN TÍCH CÁC ĐIỀU KIỆN TỰ NHIÊN, KINH TẾ - KỸ THUẬT, LỰA CHỌN ĐỊA ĐIỂM", 1, None),
    ("5.1.", "Phân tích điều kiện tự nhiên, kinh tế - kỹ thuật", 2, "5.1"),
    ("5.2.", "Lựa chọn địa điểm đầu tư", 2, "5.2"),
    # 6
    ("6.", "PHÂN TÍCH, LỰA CHỌN PHƯƠNG ÁN CÔNG NGHỆ, KỸ THUẬT, THIẾT BỊ", 1, None),
    ("6.1.", "Danh mục quy chuẩn, tiêu chuẩn chủ yếu áp dụng", 2, None),
    ("6.1.1.", "Tiêu chuẩn kỹ thuật ứng dụng CNTT trong cơ quan nhà nước", 3, "6.1.1"),
    ("6.1.2.", "Tiêu chuẩn về mạng", 3, "6.1.2"),
    ("6.1.3.", "Tiêu chuẩn về an ninh, an toàn", 3, "6.1.3"),
    ("6.1.4.", "Tiêu chuẩn áp dụng về chữ ký số", 3, "6.1.4"),
    ("6.2.", "Yêu cầu chung về kỹ thuật, công nghệ", 2, None),
    ("6.2.1.", "Yêu cầu chung về hệ thống", 3, "6.2.1"),
    ("6.2.2.", "Yêu cầu bảo đảm ATTT theo cấp độ", 3, "6.2.2"),
    ("6.2.3.", "Yêu cầu chung đối với hệ thống phần mềm", 3, "6.2.3"),
    ("6.2.4.", "Yêu cầu tính sẵn sàng IPv6", 3, "6.2.4"),
    ("6.3.", "Tiêu chí lựa chọn giải pháp công nghệ, kỹ thuật, thiết bị", 2, "6.3"),
    ("6.4.", "Phân tích lựa chọn phương án công nghệ, kỹ thuật, thiết bị", 2, None),
    ("6.4.1.", "Công nghệ mạng", 3, "6.4.1"),
    ("6.4.2.", "Công nghệ tường lửa (mạng + ứng dụng)", 3, "6.4.2"),
    ("6.4.3.", "Công nghệ lưu trữ và sao lưu (DAS/NAS/SAN)", 3, "6.4.3"),
    ("6.4.4.", "Công nghệ máy chủ (rack/blade/HCI)", 3, "6.4.4"),
    ("6.4.5.", "Công nghệ ảo hoá", 3, "6.4.5"),
    ("6.4.6.", "Công nghệ quản lý sự kiện và bảo mật thông tin (SIEM)", 3, "6.4.6"),
    ("6.4.7.", "Công nghệ giám sát mạng", 3, "6.4.7"),
    ("6.5.", "Phân tích lựa chọn giải pháp thiết kế phần mềm nội bộ", 2, None),
    ("6.5.1.", "Đề xuất ngôn ngữ lập trình", 3, "6.5.1"),
    ("6.5.2.", "Đề xuất thư viện phát triển giao diện", 3, "6.5.2"),
    ("6.5.3.", "Đề xuất mô hình phát triển phần mềm", 3, "6.5.3"),
    ("6.5.4.", "Phân tích lựa chọn hệ quản trị CSDL", 3, "6.5.4"),
    ("6.5.5.", "Phân tích lựa chọn kiến trúc phần mềm", 3, "6.5.5"),
    ("6.6.", "Phương án lựa chọn phần mềm thương mại", 2, "6.6"),
    ("6.7.", "Cơ chế phục hồi và phương án duy trì tính liên tục của hệ thống", 2, "6.7"),
    # 7
    ("7.", "THIẾT KẾ CÁC MÔ HÌNH KIẾN TRÚC CỦA HỆ THỐNG", 1, None),
    ("7.1.", "Mô hình kiến trúc tổng thể hệ thống", 2, "7.1"),
    ("7.2.", "Mô hình kiến trúc nghiệp vụ", 2, "7.2"),
    ("7.3.", "Mô hình logic hạ tầng CNTT", 2, None),
    ("7.3.1.", "Thuyết minh mô hình logic", 3, "7.3.1"),
    ("7.4.", "Mô hình vật lý hạ tầng CNTT", 2, None),
    ("7.4.1.", "Mô hình vật lý vùng trong", 3, "7.4.1"),
    ("7.4.2.", "Mô hình vật lý vùng ngoài", 3, "7.4.2"),
    ("7.4.3.", "Thuyết minh mô hình vật lý", 3, "7.4.3"),
    # 8
    ("8.", "THIẾT KẾ CƠ SỞ", 1, None),
    ("8.1.", "Phân tích tính toán định cỡ", 2, None),
    ("8.1.1.", "Nguyên tắc định cỡ", 3, "8.1.1"),
    ("8.1.2.", "Phương pháp sizing và lựa chọn", 3, "8.1.2"),
    ("8.1.3.", "Tính toán định cỡ máy chủ", 3, "8.1.3"),
    ("8.1.4.", "Tính toán định cỡ hệ thống lưu trữ", 3, "8.1.4"),
    ("8.1.5.", "Tính toán định cỡ hệ thống sao lưu", 3, "8.1.5"),
    ("8.1.6.", "Tính toán định cỡ phần mềm thương mại, hạ tầng mạng và ATTT", 3, "8.1.6"),
    ("8.1.7.", "Đề xuất đường truyền kết nối", 3, "8.1.7"),
    ("8.2.", "Thiết kế cơ sở hạ tầng phần cứng và phần mềm thương mại", 2, "8.2"),
    ("8.3.", "Thiết kế cơ sở hệ thống phần mềm nội bộ", 2, "8.3"),
    ("8.4.", "Hỗ trợ vận hành trước khi nghiệm thu", 2, None),
    ("8.4.1.", "Trực vận hành và xử lý sự cố phát sinh", 3, "8.4.1"),
    ("8.4.2.", "Trực tổng đài hỗ trợ từ xa", 3, "8.4.2"),
    ("8.5.", "Công tác đào tạo", 2, None),
    ("8.5.1.", "Mục tiêu đào tạo", 3, "8.5.1"),
    ("8.5.2.", "Nội dung và đối tượng đào tạo", 3, "8.5.2"),
    ("8.5.3.", "Yêu cầu đào tạo", 3, "8.5.3"),
    ("8.5.4.", "Yêu cầu công tác chuẩn bị cho đào tạo", 3, "8.5.4"),
    ("8.6.", "Khối lượng lắp đặt, cài đặt thiết bị và các công tác liên quan", 2, None),
    ("8.6.1.", "Khối lượng sơ bộ thiết bị và TSKT chi tiết", 3, "8.6.1"),
    ("8.6.2.", "Khối lượng công tác lắp đặt, cài đặt", 3, "8.6.2"),
    # 9
    ("9.", "THUYẾT MINH ĐỀ XUẤT CẤP ĐỘ AN TOÀN THÔNG TIN CỦA HỆ THỐNG", 1, None),
    ("9.1.", "Yêu cầu an toàn cơ bản đối với hệ thống thông tin", 2, "9.1"),
    ("9.2.", "Yêu cầu an toàn cơ bản đối với phần mềm nội bộ", 2, "9.2"),
    # 10
    ("10.", "PHƯƠNG ÁN TỔ CHỨC QUẢN LÝ, KHAI THÁC, SỬ DỤNG", 1, None),
    ("10.1.", "Phương pháp quản lý dự án", 2, None),
    ("10.1.1.", "Quản lý tiến độ dự án ứng dụng CNTT", 3, "10.1.1"),
    ("10.1.2.", "Quản lý thay đổi", 3, "10.1.2"),
    ("10.1.3.", "Quản lý chất lượng dự án", 3, "10.1.3"),
    ("10.2.", "Phương án khai thác, vận hành dự án", 2, None),
    ("10.2.1.", "Chuẩn bị về nguồn lực CNTT", 3, "10.2.1"),
    ("10.2.2.", "Chuẩn bị về cơ chế, chính sách", 3, "10.2.2"),
    ("10.2.3.", "Tổ chức hỗ trợ người sử dụng", 3, "10.2.3"),
    ("10.2.4.", "Tổ chức quản trị và vận hành hệ thống", 3, "10.2.4"),
    ("10.2.5.", "Phương án triển khai giám sát theo Luật An ninh mạng", 3, "10.2.5"),
    # 11
    ("11.", "ĐIỀU KIỆN CUNG CẤP VẬT TƯ, THIẾT BỊ, DỊCH VỤ HẠ TẦNG KỸ THUẬT, PCCC, AN NINH QUỐC PHÒNG", 1, None),
    ("11.1.", "Điều kiện cung cấp vật tư thiết bị, dịch vụ, hạ tầng kỹ thuật", 2, "11.1"),
    ("11.2.", "Điều kiện phòng chống cháy nổ, an toàn vận hành và bảo đảm an ninh", 2, None),
    ("11.2.1.", "Yêu cầu trong khu vực thi công", 3, "11.2.1"),
    ("11.2.2.", "Yêu cầu đối với cán bộ, công nhân thi công", 3, "11.2.2"),
    ("11.2.3.", "An toàn cháy nổ khi vận hành hệ thống", 3, "11.2.3"),
    ("11.2.4.", "Quy trình giải quyết sự cố", 3, "11.2.4"),
    ("11.2.5.", "Đảm bảo an toàn, bảo mật hệ thống", 3, "11.2.5"),
    ("11.3.", "Trách nhiệm khác", 2, "11.3"),
    # 12
    ("12.", "ĐÁNH GIÁ TÁC ĐỘNG VÀ GIẢI PHÁP BẢO VỆ MÔI TRƯỜNG", 1, "12"),
    # 13
    ("13.", "DỰ KIẾN TIẾN ĐỘ THỰC HIỆN DỰ ÁN", 1, "13"),
    # 14
    ("14.", "XÁC ĐỊNH TỔNG MỨC ĐẦU TƯ, CƠ CẤU NGUỒN VỐN", 1, None),
    ("14.1.", "Căn cứ lập tổng mức đầu tư", 2, "14.1"),
    ("14.2.", "Tổng hợp tổng mức đầu tư", 2, "14.2"),
    ("14.3.", "Cơ cấu nguồn vốn", 2, "14.3"),
    # 15
    ("15.", "XÁC ĐỊNH CHI PHÍ VẬN HÀNH, BẢO DƯỠNG, DUY TU, SỬA CHỮA LỚN", 1, None),
    ("15.1.", "Bảo hành và hỗ trợ kỹ thuật trong thời gian bảo hành", 2, None),
    ("15.1.1.", "Phạm vi bảo hành phần cứng", 3, "15.1.1"),
    ("15.1.2.", "Phạm vi bảo hành phần mềm thương mại", 3, "15.1.2"),
    ("15.1.3.", "Phạm vi bảo hành phần mềm ứng dụng", 3, "15.1.3"),
    ("15.1.4.", "Chi phí bảo hành", 3, "15.1.4"),
    ("15.2.", "Sơ bộ chi phí vận hành, duy tu, bảo dưỡng, sửa chữa lớn", 2, None),
    ("15.2.1.", "Dịch vụ bảo trì phần mềm ứng dụng", 3, "15.2.1"),
    ("15.2.2.", "Dịch vụ bảo trì phần mềm thương mại", 3, "15.2.2"),
    ("15.2.3.", "Dịch vụ bảo trì tổng thể tại Trung tâm dữ liệu", 3, "15.2.3"),
    # 16
    ("16.", "TỔ CHỨC QUẢN LÝ DỰ ÁN", 1, None),
    ("16.1.", "Xác định chủ đầu tư", 2, "16.1"),
    ("16.2.", "Tổ chức quản lý dự án", 2, None),
    ("16.2.1.", "Kiến nghị áp dụng hình thức quản lý dự án", 3, "16.2.1"),
    ("16.2.2.", "Nội dung công việc quản lý dự án", 3, "16.2.2"),
    ("16.3.", "Mối quan hệ và trách nhiệm của các đơn vị liên quan", 2, None),
    ("16.3.1.", "Trách nhiệm của cấp Quyết định đầu tư", 3, "16.3.1"),
    ("16.3.2.", "Trách nhiệm của Chủ đầu tư", 3, "16.3.2"),
    ("16.3.3.", "Trách nhiệm của Ban Quản lý dự án", 3, "16.3.3"),
    ("16.3.4.", "Trách nhiệm của các đơn vị thụ hưởng", 3, "16.3.4"),
    ("16.3.5.", "Trách nhiệm của Nhà thầu", 3, "16.3.5"),
    ("16.3.6.", "Trách nhiệm của các đơn vị liên quan khác", 3, "16.3.6"),
    # 17
    ("17.", "PHÂN TÍCH HIỆU QUẢ ĐẦU TƯ", 1, None),
    ("17.1.", "Hiệu quả về kinh tế, xã hội", 2, "17.1"),
    ("17.2.", "Hiệu quả về an ninh - quốc phòng", 2, "17.2"),
    # 18
    ("18.", "CÁC YẾU TỐ ĐẢM BẢO THÀNH CÔNG DỰ ÁN", 1, None),
    ("18.1.", "Phân tích rủi ro", 2, "18.1"),
    ("18.2.", "Các yếu tố đảm bảo thành công", 2, "18.2"),
    # 19
    ("19.", "KẾT LUẬN VÀ KIẾN NGHỊ", 1, "19"),
    # PL
    ("PL.", "PHỤ LỤC", 1, None),
    ("PL.1.", "Hiện trạng mặt bằng Trung tâm dữ liệu", 2, "pl.1"),
    ("PL.2.", "Sơ đồ nguyên lý triển khai hệ thống mạng", 2, "pl.2"),
    ("PL.3.", "Sơ đồ liên thông với các hệ thống liên quan", 2, "pl.3"),
]

# Diagram fields placed under their owning sections
DIAGRAMS_AT: dict[str, str] = {
    "7.1": "overall_architecture_diagram",
    "7.2": "business_architecture_diagram",
    "7.3.1": "logical_infra_diagram",
    "7.4.1": "physical_infra_inner_diagram",
    "7.4.2": "physical_infra_outer_diagram",
    "pl.1": "datacenter_layout_diagram",
    "pl.2": "network_topology_diagram",
    "pl.3": "integration_topology_diagram",
}


def _set_page_margins(section):
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_field(paragraph, instr_text: str):
    """Append a Word field code (e.g. PAGE, NUMPAGES) to paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _set_footer_page_number(section):
    """Set footer to: 'Trang {PAGE} / {NUMPAGES}' centered."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # clear default
    for r in list(p.runs):
        r.text = ""
    p.add_run("Trang ")
    _add_field(p, " PAGE ")
    p.add_run(" / ")
    _add_field(p, " NUMPAGES ")


def _set_header_running_title(section):
    """Set header to project display name (italic, right-aligned)."""
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in list(p.runs):
        r.text = ""
    run = p.add_run("BÁO CÁO NGHIÊN CỨU KHẢ THI — {{ project.display_name }}")
    run.italic = True
    run.font.size = Pt(10)


def _force_font_eastasia(style, font_name="Times New Roman"):
    """Force eastAsia font binding so VN diacritics render correctly."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def _style_heading(doc, level: int, *, size_pt: int, bold: bool, upper: bool = False):
    """Apply NĐ 30/2020 style to Heading{level}."""
    name = f"Heading {level}"
    if name not in doc.styles:
        return
    s = doc.styles[name]
    s.font.name = "Times New Roman"
    s.font.size = Pt(size_pt)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after = Pt(6)
    _force_font_eastasia(s)


def main():
    doc = Document()
    for s in doc.sections:
        _set_page_margins(s)
        _set_footer_page_number(s)
        _set_header_running_title(s)

    # ── NĐ 30/2020 Default style: Times New Roman 13, line spacing 1.5, justify ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    _force_font_eastasia(style)

    # Heading hierarchy theo NĐ 30/2020:
    # H1 (chương): bold, 14pt, UPPERCASE
    # H2 (mục):     bold, 13pt
    # H3 (tiểu):    bold, 13pt (không in nghiêng)
    _style_heading(doc, 1, size_pt=14, bold=True, upper=True)
    _style_heading(doc, 2, size_pt=13, bold=True)
    _style_heading(doc, 3, size_pt=13, bold=True)
    _style_heading(doc, 4, size_pt=13, bold=False)
    _style_heading(doc, 5, size_pt=13, bold=False)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO NGHIÊN CỨU KHẢ THI")
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DỰ ÁN: {{ project.display_name }}")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Chủ đầu tư: {{ project.client }}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Đơn vị tư vấn: {{ dev_unit }}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Ngày: {{ meta.today }} — Phiên bản: {{ meta.version }}")

    _add_page_break(doc)

    # TOC field
    p = doc.add_paragraph()
    r = p.add_run("MỤC LỤC")
    r.bold = True
    r.font.size = Pt(14)
    p_toc = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p_toc.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

    _add_page_break(doc)

    # Body
    for num, title, level, key in OUTLINE:
        # Heading — H1 chương UPPERCASE (NĐ 30/2020)
        if level == 1:
            h = doc.add_heading(level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = h.add_run(f"{num} {title.upper()}")
            r.bold = True
        elif level == 2:
            h = doc.add_heading(level=2)
            r = h.add_run(f"{num} {title}")
            r.bold = True
        elif level == 3:
            h = doc.add_heading(level=3)
            r = h.add_run(f"{num} {title}")
            r.bold = True
        else:
            h = doc.add_heading(level=min(level, 4))
            h.add_run(f"{num} {title}")

        # Diagram (if any) goes BEFORE prose for that section
        if key and key in DIAGRAMS_AT:
            field = DIAGRAMS_AT[key]
            # Conditional inline image via docxtpl
            p = doc.add_paragraph()
            p.add_run("{%p if nckt." + field + "_image %}")
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run("{{ nckt." + field + "_image }}")
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = p_cap.add_run(f"Hình. Sơ đồ §{key}")
            cap.italic = True
            p2 = doc.add_paragraph()
            p2.add_run("{%p endif %}")

        # Content placeholder
        if key:
            p = doc.add_paragraph()
            p.add_run("{{ nckt.sections.get('" + key + "', '[CẦN BỔ SUNG: §" + key + "]') }}")

    # §14.2 — Investment summary table loop (rendered AS APPENDIX after main body
    # so Jinja parser does not get confused with nested if/for in headings)
    _add_page_break(doc)
    h = doc.add_heading(level=2)
    h.add_run("Bảng tổng hợp tổng mức đầu tư (§14.2)")
    table = doc.add_table(rows=4, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Hạng mục"
    hdr[2].text = "Đơn vị"
    hdr[3].text = "Số lượng"
    hdr[4].text = "Đơn giá"
    hdr[5].text = "Thành tiền"
    hdr[6].text = "Ghi chú"
    table.rows[1].cells[0].text = "{%tr for r in nckt.investment_summary %}"
    body = table.rows[2].cells
    body[0].text = "{{ r.stt }}"
    body[1].text = "{{ r.item }}"
    body[2].text = "{{ r.unit }}"
    body[3].text = "{{ r.qty }}"
    body[4].text = "{{ r.unit_price }}"
    body[5].text = "{{ r.amount }}"
    body[6].text = "{{ r.note }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"

    # §18.1 — Risk matrix table loop
    h = doc.add_heading(level=2)
    h.add_run("Ma trận rủi ro (§18.1)")
    table = doc.add_table(rows=4, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Rủi ro"
    hdr[2].text = "Xác suất"
    hdr[3].text = "Tác động"
    hdr[4].text = "Mức độ"
    hdr[5].text = "Biện pháp giảm thiểu"
    table.rows[1].cells[0].text = "{%tr for r in nckt.risk_matrix %}"
    body = table.rows[2].cells
    body[0].text = "{{ r.stt }}"
    body[1].text = "{{ r.risk }}"
    body[2].text = "{{ r.probability }}"
    body[3].text = "{{ r.impact }}"
    body[4].text = "{{ r.level }}"
    body[5].text = "{{ r.mitigation }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK -> {OUT}  ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
