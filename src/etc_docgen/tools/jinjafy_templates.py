#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
jinjafy_templates.py — Fork ETC templates by adding Jinja2 (docxtpl) tags.

Takes original ETC template, adds:
  - {{ variable }} tags for placeholder replacements
  - {%tr for %} / {%tr endfor %} in table rows for dynamic tables
  - {%p for %} / {%p endfor %} in paragraphs for content loops
  - {% if %} / {% endif %} for conditionals

Outputs to templates/ with the standard filename (replaces pre-Jinja version).

IMPORTANT: this script is the ONLY way templates get Jinja tags added.
Don't hand-edit the .docx — re-run this script if template structure changes.

Usage:
  python jinjafy_templates.py hdsd  --source ../../../Downloads/huong-dan-su-dung.docx
  python jinjafy_templates.py tkkt  --source templates/huong-dan-su-dung.docx   # clones HDSD
  python jinjafy_templates.py tkcs  --source ../../../Downloads/Telegram\\ Desktop/TKCS_v1.2.docx
"""
from __future__ import annotations
import argparse
import re
import shutil
import sys
from pathlib import Path
from typing import Callable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WNS_QN = f"{{{WNS}}}"


# ─────────────────────────── Helpers ───────────────────────────

def para_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.findall(f".//{WNS_QN}t"))


def para_style(p_elem) -> str:
    ps = p_elem.find(f".//{WNS_QN}pStyle")
    return ps.get(f"{WNS_QN}val", "") if ps is not None else ""


def set_para_text(p_elem, new_text: str):
    """Rebuild paragraph runs with single run = new_text, keep first run formatting."""
    runs = p_elem.findall(f".//{WNS_QN}r")
    if runs:
        first_run = runs[0]
        for t in first_run.findall(f"{WNS_QN}t"):
            first_run.remove(t)
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_t.set(qn("xml:space"), "preserve")
        first_run.append(new_t)
        for r in runs[1:]:
            for t in r.findall(f"{WNS_QN}t"):
                t.text = ""
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p_elem.append(r)


def walk_all_paragraphs(doc):
    """Yield every paragraph including headers, footers, table cells."""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            yield from ncell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header,
                     section.footer, section.first_page_footer, section.even_page_footer):
            if part is None:
                continue
            yield from part.paragraphs
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def replace_everywhere(doc, replacements: list[tuple[str, str]]) -> int:
    """Walk all paragraphs, apply split-run-safe string replacements."""
    count = 0
    for p in walk_all_paragraphs(doc):
        full = "".join(r.text or "" for r in p.runs)
        if not full:
            continue
        new = full
        for old, val in replacements:
            if old in new:
                new = new.replace(old, val)
        if new != full and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
            count += 1
    return count


def add_paragraph_after(target_elem, text: str, style: str | None = None,
                        bold: bool = False, italic: bool = False):
    """Insert new paragraph right after target_elem in XML tree."""
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        new_p.append(pPr)
    new_r = OxmlElement("w:r")
    if bold or italic:
        new_rPr = OxmlElement("w:rPr")
        if bold:
            new_rPr.append(OxmlElement("w:b"))
        if italic:
            new_rPr.append(OxmlElement("w:i"))
        new_r.append(new_rPr)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    new_p.append(new_r)
    target_elem.addnext(new_p)
    return new_p


def clear_after_heading(doc, heading_text: str) -> int:
    """Delete all body elements after the heading matching text."""
    body = doc.element.body
    children = list(body)
    needle = heading_text.upper()
    marker_idx = None
    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag != "p":
            continue
        style = para_style(ch).upper()
        text = para_text(ch).upper()
        if "HEADING" in style and needle in text:
            marker_idx = i
            break
    if marker_idx is None:
        return 0
    removed = 0
    for ch in children[marker_idx + 1:]:
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag == "sectPr":
            continue
        body.remove(ch)
        removed += 1
    return removed


# ─────────────────────────── Template 1: HDSD ───────────────────────────

def jinjafy_hdsd(source: Path, dest: Path, today: str = "{{ meta.today }}"):
    """Fork ETC huong-dan-su-dung.docx with Jinja2 tags."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, dest)
    doc = Document(dest)

    # 1. Cover + header/footer placeholders
    replacements = [
        ("<Tên dự án>", "{{ project.display_name }}"),
        ("HỆ THỐNG QUẢN LÝ", "{{ project.display_name|upper }}"),
        ("CÔNG TY CỔ PHẦN HỆ THỐNG CÔNG NGHỆ ETC", "{{ dev_unit|upper }}"),
        ("<Họ tên>", "[CẦN BỔ SUNG: Tên người ký]"),
        ("<Chức danh>", "[CẦN BỔ SUNG: Chức danh]"),
        ("<Màn hình 1>", "[Minh họa màn hình]"),
        ("Ban hành: dd/mm/yyyy", "Ban hành: {{ meta.today }}"),
    ]
    n = replace_everywhere(doc, replacements)
    print(f"  Cover placeholders: {n} paragraphs updated")

    # 2. Change history (T[1] row 1) — fill with Jinja
    if len(doc.tables) > 1:
        t = doc.tables[1]
        if len(t.rows) >= 2:
            row = t.rows[1]
            cells = row.cells
            if len(cells) >= 5:
                cells[0].text = "{{ meta.today }}"
                cells[1].text = "Toàn bộ tài liệu"
                cells[2].text = "A"
                cells[3].text = "Tạo mới tài liệu"
                cells[4].text = "{{ meta.version or '1.0' }}"
                print(f"  T[1] change history: row 1 Jinja-fied")

    # 3. Abbreviations table T[4] — replace rows with Jinja loop
    if len(doc.tables) > 4:
        jinjafy_loop_table(doc.tables[4], "overview.terms",
                           ["short", "full", "explanation"])
        print(f"  T[4] abbreviations: loop added")

    # 4. Related docs T[5] — same pattern
    if len(doc.tables) > 5:
        jinjafy_loop_table(doc.tables[5], "overview.references",
                           ["stt", "name", "ref"])
        print(f"  T[5] related docs: loop added")

    # 5. Section I fills (Mục đích tài liệu, Phạm vi tài liệu)
    fill_section_i_after_heading(doc, "Mục đích tài liệu",
                                 "{{ overview.purpose }}")
    fill_section_i_after_heading(doc, "Phạm vi tài liệu",
                                 "{{ overview.scope }}")
    print(f"  Section I: purpose + scope replaced with Jinja expressions")

    # 6. Section II — clear after NỘI DUNG heading, append Jinja layout
    removed = clear_after_heading(doc, "NỘI DUNG")
    print(f"  Cleared {removed} elements after NỘI DUNG")
    append_hdsd_section_ii(doc)
    print(f"  Appended Section II Jinja layout (~50 paragraphs)")

    doc.save(dest)
    size_kb = dest.stat().st_size // 1024
    print(f"  Saved {dest} ({size_kb} KB)")


def jinjafy_loop_table(table, loop_var: str, columns: list[str]):
    """Wrap table data rows in {%tr for item %} / {%tr endfor %} — 3-row pattern.

    docxtpl requires {%tr for %} and {%tr endfor %} to be ALONE in their own rows
    (separate from data rows). Structure after this function:
      Row 0: header (untouched)
      Row 1: {%tr for item in loop_var %} alone in cell 0, other cells empty — REMOVED at render
      Row 2: data row with {{ }} expressions — REPEATED for each item
      Row 3: {%tr endfor %} alone in cell 0 — REMOVED at render
    """
    # Delete all data rows except header
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    singular = loop_var.split(".")[-1].rstrip("s")  # "terms" → "term"
    n_cols = len(table.columns)

    # Row 1: for directive alone
    for_row = table.add_row()
    for_row.cells[0].text = f"{{%tr for {singular} in {loop_var} %}}"
    for c in for_row.cells[1:]:
        c.text = ""

    # Row 2: data row with {{ }} substitutions
    data_row = table.add_row()
    for i in range(min(n_cols, len(columns))):
        col_key = columns[i]
        if col_key == "@index":
            data_row.cells[i].text = "{{ loop.index }}"
        else:
            data_row.cells[i].text = f"{{{{ {singular}.{col_key} }}}}"

    # Row 3: endfor directive alone
    end_row = table.add_row()
    end_row.cells[0].text = "{%tr endfor %}"
    for c in end_row.cells[1:]:
        c.text = ""


def fill_section_i_after_heading(doc, heading_text: str, jinja_expr: str):
    """Find heading, replace content paragraphs between it and next heading
    with a single Jinja expression paragraph."""
    body = doc.element.body
    children = list(body)
    needle = heading_text.upper()

    # Find heading
    heading_idx = None
    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag != "p":
            continue
        if "HEADING" in para_style(ch).upper() and needle in para_text(ch).upper():
            heading_idx = i
            break
    if heading_idx is None:
        return False

    # Find next heading (end boundary)
    end_idx = len(children)
    for j in range(heading_idx + 1, len(children)):
        ch = children[j]
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag == "p" and "HEADING" in para_style(ch).upper():
            end_idx = j
            break

    # Delete everything between heading and next heading
    for ch in children[heading_idx + 1:end_idx]:
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag in ("p", "tbl"):
            body.remove(ch)

    # Insert single Jinja paragraph with ETC_Content style
    add_paragraph_after(children[heading_idx], jinja_expr, style="ETC_Content")
    return True


def append_hdsd_section_ii(doc):
    """Append Jinja2 template block for Section II content after NỘI DUNG.

    Design choice: avoid nested {%p for %} + {%tr for %} — docxtpl/Jinja
    can't parse that mix. Instead use FLAT tables (no outer p-loop wrapping
    a table) and pure paragraph nesting for user manual.

    Requires render_docx.py to pre-compute `all_features` from services[].
    """
    def H(text, level):
        style = f"A_Heading {level}" if level > 1 else "A_HEADING 1"
        return doc.add_paragraph(text, style=style)

    def P(text, style="ETC_Content", bold=False, italic=False):
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return p

    # ─ 2.1 Giới thiệu chung ─
    H("Giới thiệu chung", 2)
    H("Tổng quan chương trình", 3)
    P("{{ architecture.system_overview or overview.system_description }}")
    H("Các nội dung khác", 3)
    P("{{ overview.conventions }}")

    # ─ 2.2 Giới thiệu các chức năng (FLAT combined catalog) ─
    H("Giới thiệu các chức năng", 2)
    H("Danh mục chức năng toàn hệ thống", 3)
    tbl = doc.add_table(rows=4, cols=5)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    hdr = tbl.rows[0].cells
    for i, label in enumerate(["STT", "Phân hệ", "Chức năng", "Mô tả", "Đối tượng"]):
        hdr[i].text = label
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    # Row 1: for-directive alone
    tbl.rows[1].cells[0].text = "{%tr for feat in all_features %}"
    # Row 2: data row — repeated per feature
    data = tbl.rows[2].cells
    data[0].text = "{{ loop.index }}"
    data[1].text = "{{ feat.service_name }}"
    data[2].text = "{{ feat.name }}"
    data[3].text = "{{ feat.description }}"
    data[4].text = "{{ feat.actors|join(', ') }}"
    # Row 3: endfor alone
    tbl.rows[3].cells[0].text = "{%tr endfor %}"

    # ─ 2.3 Hướng dẫn sử dụng các chức năng (PARAGRAPH-ONLY nesting) ─
    H("Hướng dẫn sử dụng các chức năng hệ thống", 2)
    P("{%p for service in services %}")
    H("{{ service.display_name }}", 3)
    P("{%p for feat in service.features %}")
    H("{{ feat.name }}", 4)
    P("{{ feat.description }}")
    P("{%p if feat.preconditions %}")
    P("Điều kiện tiên quyết: {{ feat.preconditions }}", italic=True)
    P("{%p endif %}")

    # UI elements as BULLETED LIST (paragraphs, not table — avoids nesting clash)
    P("{%p if feat.ui_elements %}")
    P("Các thành phần trên màn hình:", bold=True)
    P("{%p for elem in feat.ui_elements %}")
    P("• {{ elem.label }} ({{ elem.type }}) — {{ elem.rules or '' }}")
    P("{%p endfor %}")
    P("{%p endif %}")

    # Steps as paragraphs (each step: action + image + expected)
    P("{%p for step in feat.steps %}")
    P("Bước {{ step.no }}: {{ step.action }}", bold=True)
    P("{%p if step.screenshot_image %}")
    P("{{ step.screenshot_image }}")
    P("{%p endif %}")
    P("{%p if step.expected %}")
    P("→ Kết quả: {{ step.expected }}")
    P("{%p endif %}")
    P("{%p endfor %}")

    # Error cases (bulleted paragraphs — avoid table in nest)
    P("{%p if feat.error_cases %}")
    P("Các trường hợp lỗi:", bold=True)
    P("{%p for err in feat.error_cases %}")
    P("• Bước {{ err.trigger_step }}: {{ err.condition }} → {{ err.message }}")
    P("{%p endfor %}")
    P("{%p endif %}")

    P("{%p endfor %}")  # end features
    P("{%p endfor %}")  # end services

    # ─ 2.4 Các vấn đề thường gặp (FLAT table, 3-row pattern) ─
    H("Các vấn đề thường gặp khi sử dụng", 2)
    tbl3 = doc.add_table(rows=4, cols=4)
    try:
        tbl3.style = "Table Grid"
    except KeyError:
        pass
    hdr3 = tbl3.rows[0].cells
    for i, label in enumerate(["STT", "Tình huống", "Nguyên nhân", "Cách xử lý"]):
        hdr3[i].text = label
        for r in hdr3[i].paragraphs[0].runs:
            r.bold = True
    tbl3.rows[1].cells[0].text = "{%tr for item in troubleshooting %}"
    dr = tbl3.rows[2].cells
    dr[0].text = "{{ loop.index }}"
    dr[1].text = "{{ item.situation }}"
    dr[2].text = "{{ item.cause }}"
    dr[3].text = "{{ item.resolution }}"
    tbl3.rows[3].cells[0].text = "{%tr endfor %}"


# ─────────────────────────── Template 2: TKKT ───────────────────────────

def jinjafy_tkkt(source: Path, dest: Path):
    """TKKT uses HDSD cover + signing pages, but different Section II content."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, dest)
    doc = Document(dest)

    # 1. Cover + header/footer (same as HDSD, plus replace HDSD title → TKKT)
    replacements = [
        ("<Tên dự án>", "{{ project.display_name }}"),
        ("HỆ THỐNG QUẢN LÝ", "{{ project.display_name|upper }}"),
        ("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG", "TÀI LIỆU THIẾT KẾ KIẾN TRÚC"),
        ("Tài liệu hướng dẫn sử dụng", "Tài liệu thiết kế kiến trúc"),
        ("CÔNG TY CỔ PHẦN HỆ THỐNG CÔNG NGHỆ ETC", "{{ dev_unit|upper }}"),
        ("<Họ tên>", "[CẦN BỔ SUNG: Tên người ký]"),
        ("<Chức danh>", "[CẦN BỔ SUNG: Chức danh]"),
        ("<Màn hình 1>", ""),
        ("Ban hành: dd/mm/yyyy", "Ban hành: {{ meta.today }}"),
    ]
    replace_everywhere(doc, replacements)

    # 2. Change history
    if len(doc.tables) > 1:
        t = doc.tables[1]
        if len(t.rows) >= 2:
            row = t.rows[1]
            cells = row.cells
            if len(cells) >= 5:
                cells[0].text = "{{ meta.today }}"
                cells[1].text = "Toàn bộ tài liệu"
                cells[2].text = "A"
                cells[3].text = "Tạo mới tài liệu thiết kế kiến trúc"
                cells[4].text = "{{ meta.version or '1.0' }}"

    # 3. Abbreviations + Related docs (same as HDSD)
    if len(doc.tables) > 4:
        jinjafy_loop_table(doc.tables[4], "overview.terms",
                           ["short", "full", "explanation"])
    if len(doc.tables) > 5:
        jinjafy_loop_table(doc.tables[5], "overview.references",
                           ["stt", "name", "ref"])

    # 4. Section I
    fill_section_i_after_heading(doc, "Mục đích tài liệu",
                                 "{{ architecture.purpose }}")
    fill_section_i_after_heading(doc, "Phạm vi tài liệu",
                                 "{{ architecture.scope }}")

    # 5. Section II — architecture-specific layout
    clear_after_heading(doc, "NỘI DUNG")
    append_tkkt_section_ii(doc)

    doc.save(dest)
    print(f"  Saved {dest} ({dest.stat().st_size // 1024} KB)")


def append_tkkt_section_ii(doc):
    """7 sections: overview / logic / data / integration / deployment / security / NFR."""
    def H(text, level):
        style = f"A_Heading {level}" if level > 1 else "A_HEADING 1"
        return doc.add_paragraph(text, style=style)

    def P(text, style="ETC_Content"):
        return doc.add_paragraph(text, style=style)

    def simple_table(hdr_labels, rows_jinja, loop_var, col_exprs):
        """Build table with header + 3-row Jinja loop pattern."""
        tbl = doc.add_table(rows=4, cols=len(hdr_labels))
        try:
            tbl.style = "Table Grid"
        except KeyError:
            pass
        # Header
        h = tbl.rows[0].cells
        for i, label in enumerate(hdr_labels):
            h[i].text = label
            for r in h[i].paragraphs[0].runs:
                r.bold = True
        # for-directive row
        singular = loop_var.split(".")[-1].rstrip("s")
        tbl.rows[1].cells[0].text = f"{{%tr for {singular} in {loop_var} %}}"
        # data row — use `item` variable name to match existing col_exprs which use {{ item.x }}
        # col_exprs use 'item' but loop var is `singular` — rewrite col exprs to use singular
        data = tbl.rows[2].cells
        for i, expr in enumerate(col_exprs):
            # Replace {{ item. with {{ singular.
            rewritten = expr.replace("{{ item.", f"{{{{ {singular}.")
            rewritten = rewritten.replace("item.", f"{singular}.")
            data[i].text = rewritten
        # endfor row
        tbl.rows[3].cells[0].text = "{%tr endfor %}"

    # 2.1 Tổng quan hệ thống
    H("Tổng quan hệ thống", 2)
    P("{{ architecture.system_overview }}")
    H("Phạm vi và đối tượng", 3)
    P("{{ architecture.scope_description }}")
    H("Công nghệ sử dụng", 3)
    simple_table(["Lớp", "Công nghệ", "Phiên bản", "Vai trò"],
                 None, "architecture.tech_stack",
                 ["{{ item.layer }}", "{{ item.technology }}",
                  "{{ item.version or '-' }}", "{{ item.role }}"])

    # 2.2 Kiến trúc logic
    H("Kiến trúc logic", 2)
    P("{{ architecture.logical_description }}")
    H("Danh sách thành phần", 3)
    simple_table(["STT", "Thành phần", "Loại", "Mô tả"],
                 None, "architecture.components",
                 ["{{ loop.index }}", "{{ item.name }}",
                  "{{ item.type }}", "{{ item.description }}"])
    H("Tương tác giữa các thành phần", 3)
    P("{{ architecture.interaction_description }}")

    # 2.3 Kiến trúc dữ liệu
    H("Kiến trúc dữ liệu", 2)
    P("{{ architecture.data_description }}")
    H("Mô hình dữ liệu", 3)
    simple_table(["STT", "Bảng/Collection", "Mục đích", "Loại lưu trữ"],
                 None, "architecture.data_entities",
                 ["{{ loop.index }}", "{{ item.name }}",
                  "{{ item.purpose }}", "{{ item.storage_type or 'PostgreSQL' }}"])

    # 2.4 Kiến trúc tích hợp
    H("Kiến trúc tích hợp", 2)
    P("{{ architecture.integration_description }}")
    H("Danh sách API/Interface", 3)
    simple_table(["STT", "Endpoint", "Phương thức", "Mô tả", "Xác thực"],
                 None, "architecture.apis",
                 ["{{ loop.index }}", "{{ item.path }}", "{{ item.method }}",
                  "{{ item.description }}", "{{ item.auth or 'JWT Bearer' }}"])
    H("Tích hợp hệ thống ngoài", 3)
    simple_table(["STT", "Hệ thống ngoài", "Giao thức", "Mục đích"],
                 None, "architecture.external_integrations",
                 ["{{ loop.index }}", "{{ item.system }}",
                  "{{ item.protocol }}", "{{ item.purpose }}"])

    # 2.5 Kiến trúc triển khai
    H("Kiến trúc triển khai", 2)
    P("{{ architecture.deployment_description }}")
    H("Môi trường triển khai", 3)
    simple_table(["Môi trường", "Hạ tầng", "Mục đích"],
                 None, "architecture.environments",
                 ["{{ item.name }}", "{{ item.infrastructure }}", "{{ item.purpose }}"])
    H("Các container/service", 3)
    simple_table(["STT", "Container", "Image", "Port", "Phụ thuộc"],
                 None, "architecture.containers",
                 ["{{ loop.index }}", "{{ item.name }}", "{{ item.image }}",
                  "{{ item.port or '-' }}", "{{ item.depends_on|join(', ') if item.depends_on else '-' }}"])

    # 2.6 Kiến trúc bảo mật
    H("Kiến trúc bảo mật", 2)
    P("{{ architecture.security_description }}")
    H("Xác thực và phân quyền", 3)
    P("{{ architecture.auth_description }}")
    H("Bảo vệ dữ liệu", 3)
    P("{{ architecture.data_protection or '[CẦN BỔ SUNG: encryption at rest, TLS, masking]' }}")

    # 2.7 Yêu cầu phi chức năng
    H("Yêu cầu phi chức năng", 2)
    simple_table(["Tiêu chí", "Yêu cầu", "Phương án đáp ứng"],
                 None, "architecture.nfr",
                 ["{{ item.criterion }}", "{{ item.requirement }}", "{{ item.solution }}"])


# ─────────────────────────── Template 3: TKCS ───────────────────────────

def jinjafy_tkcs(source: Path, dest: Path):
    """TKCS (based on stripped TKCS_v1.2) — minimal fills only.

    Strategy: fill Section 1 (Giới thiệu chung) with project info from Jinja,
    and prepend [CẦN BỔ SUNG] hints under each top-level heading for BA.
    """
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, dest)
    doc = Document(dest)

    # Cover replacements
    replacements = [
        ("<Tên dự án>", "{{ project.display_name }}"),
        ("CÔNG TY CỔ PHẦN HỆ THỐNG CÔNG NGHỆ ETC", "{{ dev_unit|upper }}"),
        ("<Họ tên>", "[CẦN BỔ SUNG: Tên người ký]"),
        ("<Chức danh>", "[CẦN BỔ SUNG: Chức danh]"),
    ]
    replace_everywhere(doc, replacements)

    # Section 1 fills — find H2 headings under Giới thiệu chung
    fills_section1 = [
        ("Tên dự án", "{{ project.display_name }}"),
        ("Chủ đầu tư", "{{ project.client }}"),
        ("Tổ chức tư vấn lập dự án", "{{ dev_unit }}"),
        ("Hình thức đầu tư", "[CẦN BỔ SUNG: Đầu tư mới / Nâng cấp / Thuê dịch vụ]"),
        ("Thời gian thực hiện dự án", "[CẦN BỔ SUNG: Thời gian dự kiến]"),
        ("Nguồn vốn", "[CẦN BỔ SUNG: Ngân sách nhà nước / Vốn tự có]"),
    ]
    for heading, content in fills_section1:
        replace_content_after_h2(doc, heading, content)

    # Section-level hints — inserted after each H1
    hints = [
        ("SỰ CẦN THIẾT ĐẦU TƯ",
         "[CẦN BỔ SUNG: BA viết cơ sở pháp lý, hiện trạng, sự cần thiết đầu tư]"),
        ("ĐÁNH GIÁ SỰ PHÙ HỢP VỚI QUY HOẠCH",
         "[CẦN BỔ SUNG: Thuyết minh phù hợp Khung Kiến trúc CPĐT 4.0, Kế hoạch CNTT]"),
        ("PHÂN TÍCH, LỰA CHỌN PHƯƠNG ÁN CÔNG NGHỆ",
         "{{ tkcs.technology_rationale or '[CẦN BỔ SUNG: Phân tích lựa chọn công nghệ]' }}"),
        ("THIẾT KẾ CƠ SỞ PHƯƠNG ÁN CHỌN",
         "{{ tkcs.detailed_design_summary or '[CẦN BỔ SUNG: Tham chiếu Tài liệu Thiết kế Kiến trúc]' }}"),
        ("PHƯƠNG ÁN ĐẢM BẢO AN TOÀN THÔNG TIN",
         "[CẦN BỔ SUNG: Đánh giá cấp độ ATTT theo TT 03/2017/TT-BTTTT]"),
        ("PHƯƠNG ÁN TỔ CHỨC QUẢN LÝ, KHAI THÁC",
         "[CẦN BỔ SUNG: Mô hình tổ chức vận hành, SLA]"),
        ("DỰ KIẾN TIẾN ĐỘ THỰC HIỆN",
         "[CẦN BỔ SUNG: Mốc thời gian triển khai]"),
        ("XÁC ĐỊNH TỔNG MỨC ĐẦU TƯ",
         "[CẦN BỔ SUNG: Bóc khối lượng theo TT 04/2020/TT-BTTTT]"),
        ("XÁC ĐỊNH CHI PHÍ VẬN HÀNH",
         "[CẦN BỔ SUNG: Chi phí O&M hàng năm]"),
        ("TỔ CHỨC QUẢN LÝ DỰ ÁN",
         "[CẦN BỔ SUNG: Mô hình QLDA — BQL chuyên trách / kiêm nhiệm]"),
    ]
    for heading, content in hints:
        insert_after_h1(doc, heading, content, italic=True)

    doc.save(dest)
    print(f"  Saved {dest} ({dest.stat().st_size // 1024} KB)")


def replace_content_after_h2(doc, heading_text: str, new_content: str):
    """Find H2 heading, replace the single paragraph after it with new_content."""
    body = doc.element.body
    children = list(body)
    needle = heading_text.upper()
    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag != "p":
            continue
        if "heading" in para_style(ch).lower() and needle in para_text(ch).upper():
            # Replace first following non-heading paragraph
            for j in range(i + 1, min(i + 5, len(children))):
                nxt = children[j]
                ntag = nxt.tag.split("}")[-1] if "}" in nxt.tag else nxt.tag
                if ntag != "p":
                    continue
                nstyle = para_style(nxt).lower()
                if "heading" in nstyle:
                    # No content paragraph between headings — insert one
                    add_paragraph_after(ch, new_content)
                    return
                # Replace its text
                set_para_text(nxt, new_content)
                return
            break


def insert_after_h1(doc, heading_text: str, content: str, italic: bool = False):
    """Insert paragraph right after H1 heading matching text."""
    body = doc.element.body
    needle = heading_text.upper()
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "p":
            continue
        if "heading" in para_style(child).lower() and needle in para_text(child).upper():
            add_paragraph_after(child, content, italic=italic)
            return


# ─────────────────────────── CLI ───────────────────────────

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("kind", choices=["hdsd", "tkkt", "tkcs"])
    ap.add_argument("--source", required=True, help="Source template (original ETC or cloned)")
    ap.add_argument("--dest", default=None, help="Output path; defaults to templates/<name>.docx")
    args = ap.parse_args()

    source = Path(args.source)
    if not source.exists():
        print(f"ERROR: source not found: {source}", file=sys.stderr)
        sys.exit(1)

    default_names = {
        "hdsd": "huong-dan-su-dung.docx",
        "tkkt": "thiet-ke-kien-truc.docx",
        "tkcs": "thiet-ke-co-so.docx",
    }
    dest = Path(args.dest) if args.dest else (
        Path(__file__).parent.parent / "templates" / default_names[args.kind]
    )

    print(f"Jinjafying {args.kind.upper()}: {source} → {dest}")
    if args.kind == "hdsd":
        jinjafy_hdsd(source, dest)
    elif args.kind == "tkkt":
        jinjafy_tkkt(source, dest)
    elif args.kind == "tkcs":
        jinjafy_tkcs(source, dest)


if __name__ == "__main__":
    main()
